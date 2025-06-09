from openai import OpenAI
from dotenv import load_dotenv
import os
import json
import math
import pandas as pd
pd.set_option('display.max_columns', None)
pd.set_option('display.max_rows', 100)
from difflib import SequenceMatcher
from docx import Document
from image_details_extractor import generate_product_description
from analytics_matcher import match_headline_to_keyword
import re

# Load environment variables from .env file
load_dotenv()

# Initialize the OpenAI client using your endpoint and token
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
)

# Sheets: ['Model Training', '1 New Romance Copy Generation', '2 New Products Part of MegaPDP\u200b', '3 Products for SEO Enrichment', 'Full Catalog Short Descriptions']
# Load Coach Rules
coach_doc = Document("Documents/Coach Rules.docx")
coach_rules = "\n".join([para.text for para in coach_doc.paragraphs])

# Load Kate Spade Rules
spade_doc = Document("Documents/Kate Spade Rules.docx")
spade_rules = "\n".join([para.text for para in spade_doc.paragraphs])

# Data File paths
file1 = 'Documents\POC Product Selection- Coach Outlet.xlsx'
file2 = 'Documents\POC Product Selection- Kate Spade.xlsx'

coach_sheets = pd.read_excel(file1, sheet_name=None)
spade_sheets = pd.read_excel(file2, sheet_name=None)

def similarity(a: str, b: str) -> float:
    """Return a float [0.0–1.0] for how similar two strings are."""
    return SequenceMatcher(None, a, b).ratio()

def get_tagline(product_attributes, product_description_image, analytics, company):
    raw_mega_value = product_attributes.get("Mega PDP Group Value", "")
    mega_value = str(raw_mega_value).lower() if pd.notna(raw_mega_value) else ""

    if company =="Spade":
        rules = spade_rules
        mega_unique_values = (
            set(val.lower() for val in spade_sheets["Model Training"]["Mega PDP Group Value"].dropna().unique())
            .union(
                val.lower() for val in spade_sheets["Full Catalog Short Descriptions"]["Mega PDP Group Value"].dropna().unique()
            )
        )

        prod_old_description = []
        matched_mega_values = []
        match_type = ""

        # Check for exact match
        if mega_value in mega_unique_values:
            print("Exact match")
            match_type = "Exact"
            for sheet_name in ("Model Training", "Full Catalog Short Descriptions"):
                df = spade_sheets[sheet_name]
                matches = df[df["Mega PDP Group Value"].str.lower() == mega_value]
                if not matches.empty:
                    prod_old_description.extend(matches["Short Description - en"].tolist())
                    matched_mega_values.extend(matches["Mega PDP Group Value"].tolist())
        
        # If no exact match, try similar match
        else:
            similarity_scores = [(val, similarity(mega_value, val)) for val in mega_unique_values]
            similarity_scores.sort(key=lambda x: x[1], reverse=True)
            
            if similarity_scores:
                def fetch_matches(match_vals):
                    for match_val in match_vals:
                        for sheet_name in ("Model Training", "Full Catalog Short Descriptions"):
                            df = spade_sheets[sheet_name]
                            matches = df[df["Mega PDP Group Value"].str.lower() == match_val]
                            if not matches.empty:
                                prod_old_description.extend(matches["Short Description - en"].tolist())
                                matched_mega_values.extend(matches["Mega PDP Group Value"].tolist())
                                return True  # Stop at first valid match
                    return False

                top_match_val, top_match_score = similarity_scores[0]

                if top_match_score >= 0.7:
                    print(f"Found similar match: {top_match_val} (Score: {top_match_score})")
                    match_type = "Similar threshold greater than 70%"
                    fetch_matches([top_match_val])
                else:
                    print("No strong match found, returning top 10 matches")
                    match_type = "Less than 70%"
                    top_10_match_vals = [val for val, _ in similarity_scores[:10]]
                    for val, score in similarity_scores[:10]:
                        print(f"Match: {val} (Score: {score})")
                    fetch_matches(top_10_match_vals)
        blacklisted_keywords = [
            "earned a treat", "NO.really", "s.a.l.e.", "expires", "celebrating", "psst", 
            "customer", "hello", "sale on sale", "leaving soon", "Rewarding", "surprise", 
            "elevate", "girl on the go", "hang", "you've bagged", "it's your final chance", 
            "treating you to code", "e_legance", "elegant", "hot", "vintage", "discount", 
            "attention", "you have", "glamorous", "kitsch", "lady", "splurge", "Official", 
            "we're releasing", "discover", "ob_sessed", "got to", "major bag alert", "Officially", 
            "releasing", "open immediately", "retro", "Win", "edgy", "all for you", 
            "you're getting", "order today", "utterly", "#win", "open asap", "now trending", 
            "confirm", "Announcement", "chic", "deal", "1-day", "yes", "all caps", 
            "Announcing", "officially in stock", "adorable", "(1-day special!)", "take", 
            "Lucky you", "Score", "cute", "fresh", "released", "explore", "presenting", 
            "all eyes on", "classy", "gorgeous", "markdown", "Checkout", "no joke", 
            "for you", "awesome", "hung", "Babe", "redeem", "Oooh", "get one", 
            "is sure to excite", "smile", "snack", "hey", "reserved", "make one yours", 
            "nice", "Knott", "as a thank you", "calling your name", "P_ssst", "Psst", 
            "view", "tons", "oh", "no", "earn", "just in", "flirty", "secure", 
            "hello gorgeous", "oof", "glow on", "just reduced", "sexy", "Deserve", 
            "hello", "gorgeous", "sale just dropped", "buy more", "save more", "unlock", 
            "name a more iconic", "Shop", "kind of time-sensitive", "we're confirming", 
            "offering", "treat", "duo", "Styles made to last", "must-have", "alert", 
            "compliments of us", "claim", "New you", "Enhance", "special message", 
            "you're receiving", "upgraded", "we're giving you", "One-day", "No exclusions", 
            "special feature", "just-reduced", "shipment", "hi there", "Snag", "Expires", 
            "girl", "sale confirmed", "wristlet", "Hey you", "Continue", "Leaving soon", 
            "because you rock", "you've secured", "all emojis", "Landed", "check out", 
            "It's your final chance", "the modern woman", "fashion-forward individual", "smart", 
            "PVC", "sophisticated", "modern wardrobe", "luxurious", "logo", 
            "logo embellishment", "Saffiano PVC", "the modern woman", "smart", 
            "sophisticated", "fashion-forward individual", "modern wardrobe", "sophistication", 
            "we", "you", "casual day", "flair", "causal outings", "casual", 
            "brighter days", "metal material", "sophistication", "trust us", "day party", 
            "fashion-savvy individual", "elegance", "elegant", "modern fashion", "modern","precision edge painting"
        ]
    if company =="Coach":
        rules= coach_rules
        mega_unique_values = (
            set(val.lower() for val in coach_sheets["Model Training"]["Mega PDP Group Value"].dropna().unique())
            .union(
                val.lower() for val in coach_sheets["Full Catalog Short Descriptions"]["Mega PDP Group Value"].dropna().unique()
            )
        )

        prod_old_description = []
        matched_mega_values = []
        match_type = ""

        # Check for exact match
        if mega_value in mega_unique_values:
            print("Exact match")
            match_type = "Exact"
            for sheet_name in ("Model Training", "Full Catalog Short Descriptions"):
                df = coach_sheets[sheet_name]
                matches = df[df["Mega PDP Group Value"].str.lower() == mega_value]
                if not matches.empty:
                    prod_old_description.extend(matches["Short Description - en"].tolist())
                    matched_mega_values.extend(matches["Mega PDP Group Value"].tolist())
        
        # If no exact match, try similar match
        else:
            similarity_scores = [(val, similarity(mega_value, val)) for val in mega_unique_values]
            similarity_scores.sort(key=lambda x: x[1], reverse=True)
            
            if similarity_scores:
                def fetch_matches(match_vals):
                    for match_val in match_vals:
                        for sheet_name in ("Model Training", "Full Catalog Short Descriptions"):
                            df = coach_sheets[sheet_name]
                            matches = df[df["Mega PDP Group Value"].str.lower() == match_val]
                            if not matches.empty:
                                prod_old_description.extend(matches["Short Description - en"].tolist())
                                matched_mega_values.extend(matches["Mega PDP Group Value"].tolist())
                                return True  # Stop at first valid match
                    return False

                top_match_val, top_match_score = similarity_scores[0]

                if top_match_score >= 0.7:
                    print(f"Found similar match: {top_match_val} (Score: {top_match_score})")
                    match_type = "Similar threshold greater than 70%"
                    fetch_matches([top_match_val])
                else:
                    print("No strong match found, returning top 10 matches")
                    match_type = "Less than 70%"
                    top_10_match_vals = [val for val, _ in similarity_scores[:10]]
                    for val, score in similarity_scores[:10]:
                        print(f"Match: {val} (Score: {score})")
                    fetch_matches(top_10_match_vals)
        blacklisted_keywords = [
            "inspired by", "chic", "exudes sophistication", "gen-z customer", "gen-z", 
            "aesthetic", "affordable", "ageless", "body", "chic", "coachie", "couture", 
            "craftsman", "customer", "cute", "dainty", "daintier", "darling", "deal", 
            "delightful", "designer", "discount", "disruptive", "don", "donning", 
            "easy win", "elegant", "elegance", "embellished", "enchanting", "engineered", 
            "eternal", "expressive luxury", "fabulous", "fabulousness", "fashion", 
            "fashion lover", "fashionista", "fave", "footwear", "gang", "gender-neutral", 
            "handbag", "hot", "it bag", "it girl", "it’s giving", "jet-set", "lovely", 
            "multifunctional", "must have", "new you", "obsess", "obsessed", "mindful", 
            "green", "conscious", "eco-conscious", "pioneering", "pleasing", "pretty", 
            "purse", "quiet luxury", "sale", "sassy", "savage", "sensations", "sleek", 
            "splendid", "sueded", "sustainable", "szn", "tender", "treasures", 
            "trendsetter", "turn heads", "unearth", "unveil", "unveiling", "uptown style", 
            "downtown style", "urban", "vibes", "but make it fashion", "meet", 
            "experience", "introducing", "just", "literally", "figuratively", 
            "audacious", "pvc", "PVC", "mundane", "nitty-gritties", "beauty scores", 
            "best", "boast", "booster", "statement", "promise", "declaration", "go-to", 
            "taste", "impeccable", "pretty face", "testament", "touches", "must-have", 
            "impraczcal", "unassuming", "overlook", "unusual", "friend", "flair", 
            "fierce", "efforzless", "glamour", "outing", "fashionable", "stylish", 
            "more than a pretty face", "your new best friend", "accessories collection", 
            "boasts", "modern fashion", "this is Coach Outlet's promise to you", 
            "declaration of style", "testament to your impeccable taste", 
            "finishing touches from Coach Outlet", "fashion adventures", "flair", 
            "audacious modern style", "this beauty scores high", "meziculously", 
            "captivating", "aesthetics", "simplicity and class", "dash of the unusual", 
            "crafted to fulfill", "unassuming elegance", "it's impractical to overlook", 
            "sexy", "let's talk about", "inspiration can come", "fall in love", 
            "inspiration can strike", "picture this", "picture themselves", "imagine", 
            "bio-attributed", "bio-based", "biodegradable", "bio-finished", 
            "carbon neutral", "certified b corp", "chemical recycling", "circular", 
            "closed loop", "compostable", "fair trade", "FSC", "forest stewardship", 
            "council", "mechanical recycling", "natural", "PEFC", "recyclable", 
            "upcycled", "SFI", "responsible", "synthetic", "traceable", "transparent", 
            "vegan", "zero waste"
        ]

    prod_old_description = list(set(prod_old_description))
    def remove_blacklisted_keywords(paragraphs, blacklisted_keywords):
        cleaned_paragraphs = []
        
        for paragraph in paragraphs:
            if not isinstance(paragraph, str) or not paragraph.strip():
                continue  # Skip NaN, empty strings, or non-string values
            
            cleaned_paragraph = paragraph
            for keyword in blacklisted_keywords:
                pattern = r'\b' + re.escape(keyword) + r'\b'
                cleaned_paragraph = re.sub(pattern, '', cleaned_paragraph, flags=re.IGNORECASE)
            
            cleaned_paragraph = ' '.join(cleaned_paragraph.split())
            cleaned_paragraphs.append(cleaned_paragraph)
        
        # Remove duplicates while preserving order
        seen = set()
        unique_paragraphs = [p for p in cleaned_paragraphs if not (p.lower() in seen or seen.add(p.lower()))]
        
        return unique_paragraphs


    # Apply the function to prod_old_description
    prod_old_description = remove_blacklisted_keywords(prod_old_description, blacklisted_keywords)
    prompt = [
    "Instructions:",
    "1. You are given a set of rules. Follow them exactly to generate a new tagline.",
    "2. Do NOT repeat any word in the tagline.",
    "3. Do NOT include the phrase “what fits inside.”",
    "4. Do NOT include any city-specific references.",
    "5. Do NOT mention pairing dress or attire.",
    "6. Do NOT use general phrases—be product-specific.",
    "7. Maintain a natural, authentic tone.",
    "8. Strictly avoid all blacklisted words (severe penalty for violations).",
    "9. Strictly follow the structure and the content of sample given, else you will be heavily penalized.",
    "**You must follow Instructions and rules at any cost, else you will be heavily penalized.**",
    "####",
    "Rules:",
    f"{rules}",
    "####",
    ]

    if prod_old_description != "":
        prompt.append("**Strictly follow the structure and the content of sample given below, to generate the new tagline. Keep the structure and content same to the given sample. Do not add Blacklisted words.**")
        prompt.append(f"{prod_old_description}")

    prompt += [
        "####",
        "SEO Guidance:",
        "Generate an SEO keyword list.",
        "Keyword Hierarchy (incorporate these into your tagline where natural):",
        "- **Primary Keywords** - Describing product type (example - Flap shoulder bag, Colorblocked bag, Convertible bag)",
        "- **Secondary Keywords** - Describing characteristics (example - Pebbled leather, Colorblocked leather, Classic flap silhouette, Adjustable crossbody strap, Convertible design)",
        "- **Tertiary Keywords** - Describing function - (example - Optional crossbody strap, Everyday bag, Versatile handbag)",
    ]

    if analytics != {}:
        prompt+= [
            "####",
            "Analyze the Google Analytics report below:", 
            f"{analytics}"
        ]

    if product_description_image != {}:
        prompt.extend(["####"," Below is the visual description of the image. Do take into account while framing the Tagline.",
            f"{product_description_image}"])
    
    prompt.append("####")
    prompt.append("Below are the attributes for the product:")

     # Dynamically iterate over every top-level key in the dictionary
    for key, value in product_attributes.items():
        if key in ["What Fits Inside - en","Iteration","Tech Fit - en","Primary Digital Asset URL","Non-Primary Digital Asset URL"]:
            continue
        if value == "" or (isinstance(value, float) and math.isnan(value)):
            continue
        pretty_value = json.dumps(value, indent=2, ensure_ascii=False)
        # Prepend each line of the value with two spaces so it's clear it belongs under the key
        indented_value = "\n".join([f"  {line}" for line in pretty_value.splitlines()])
        # Append a line like "Key:" followed by the indented JSON value
        prompt.append(f"- {key}:\n{indented_value}")

    prompt.extend([
        f"Before generating the tagline, Please check: Do NOT use these blacklisted words under any circumstances: {blacklisted_keywords}.",
        "Do not use you, your,etc in tagline."
    ])

    prompt.extend([
    "Format your response exactly like this (so it’s easy to parse):",
        "```json",
        "{",
        '  "editorial_tagline": "...",',
        '  "SEO Keyword 1": ["...", "...", "..."],',
        '  "SEO Keyword 2": ["...", "...", "..."],',
        '  "SEO Keyword 3": ["...", "...", "..."]',
        "}",
        "```"
    ])

    full_prompt = "\n".join(prompt)

    system_prompts = [
    "You are a world-class luxury fashion editor. Do NOT add Blacklisted Words in the tagline."
    "Instructions:",
    "1. You are given a set of rules. Follow them exactly to generate a new tagline.",
    "2. Do NOT repeat any word in the tagline.",
    "3. Do NOT include the phrase “what fits inside.”",
    "4. Do NOT include any city-specific references.",
    "5. Do NOT mention pairing dress or attire.",
    "6. Do NOT use general phrases—be product-specific.",
    "7. Maintain a natural, authentic tone.",
    "8. Strictly avoid all blacklisted words (severe penalty for violations).",
    "9. Strictly follow the structure and the content of sample given, else you will be heavily penalized.",
    "**You must follow Instructions and rules at any cost, else you will be heavily penalized.**",
    ]

    system_prompts = "\n".join(system_prompts)

    # print(full_prompt)

    response = client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": system_prompts},
                {"role": "user", "content": full_prompt}
            ],
            temperature=0.3,
            response_format={"type": "json_object"}
        )
    
    res = json.loads(response.choices[0].message.content.strip())

    found_blacklisted_keywords = []

    tagline = res["editorial_tagline"]
    for keyword in blacklisted_keywords:
        pattern = r'\b' + re.escape(keyword) + r'\b'
        if re.search(pattern, tagline, re.IGNORECASE):
            found_blacklisted_keywords.append(keyword)
        

    if found_blacklisted_keywords:
        print(f"🚨 🚨 🚨 🚨 🚨 🚨 Blacklisted keywords present: {', '.join(set(found_blacklisted_keywords))},")
        print(tagline)
    else:
        print("No blacklisted keywords found.")

    res["Old Description"] = prod_old_description
    res["Matched OLD Mega PDP Value"] = matched_mega_values
    res["Prompt"] = full_prompt
    res["Match_Type"] = match_type
    print(res["editorial_tagline"])

    return res

def process_usecase(usecase_df, brand):
    data = usecase_df.to_dict(orient='records')
    output_data = []
    
    for item in data:
        print(f"Processing {item['Item#']}")
        image = item.get("Primary Digital Asset URL", "")
        image2 = item.get("Primary Digital Asset URL", "")  # Note: This might need adjustment if a secondary image column exists
        raw_urls = f"{image}`{image2}".replace("`", ",").split(",")
        images = list(filter(None, map(str.strip, raw_urls)))
        
        if images:
            product_description_image = generate_product_description(images)
        else:
            product_description_image = {}
        
        product_name = item.get("Web Product Name - en", [])  # Adjusted to match requested column name
        if product_name:
            analytics = match_headline_to_keyword(product_name)
        else:
            analytics = {}
        
        luxury_tagline = get_tagline(item, product_description_image, analytics, brand)
        
        if isinstance(luxury_tagline, dict):
            for k, v in luxury_tagline.items():
                item[k] = v
        else:
            item["Luxury Tagline"] = luxury_tagline
        
        output_data.append(item)
    
    return pd.DataFrame(output_data)

def main():
    # Define all use cases for Coach and Spade
    coach_usecase1 = coach_sheets["1 New Romance Copy Generation"]
    coach_usecase2 = coach_sheets["2 New Products Part of MegaPDP\u200b"]
    coach_usecase3 = coach_sheets["3 Products for SEO Enrichment"]
    
    # spade_usecase1 = spade_sheets["1 New Romance Copy Generation"]
    # spade_usecase2 = spade_sheets["2 New Products Part of MegaPDP\u200b"]
    # spade_usecase3 = spade_sheets["3 Products for SEO Enrichment"]
    
    # List of use cases with their respective brands
    usecases = [
        (coach_usecase1, "Coach"),
        (coach_usecase2, "Coach"),
        (coach_usecase3, "Coach"),
        # (spade_usecase1, "Spade"),
        # (spade_usecase2, "Spade"),
        # (spade_usecase3, "Spade"),
    ]
    
    all_results = []
    
    # Process each use case and collect results
    for usecase_df, brand in usecases:
        processed_df = process_usecase(usecase_df, brand)
        all_results.append(processed_df)
    
    # Concatenate all processed DataFrames
    all_results_df = pd.concat(all_results, ignore_index=True)
    
    # Save first Excel file with all results in original format
    all_results_df.to_excel("All_Usecases_Results.xlsx", index=False, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}})
    print("Saved to All_Usecases_Results.xlsx")
    
    # Define columns for the new format Excel file
    selected_columns = [
        "Item#",
        "Web Product Name - en",
        "Mega PDP Group Value",
        "editorial_tagline",
        "SEO Keyword 1",
        "SEO Keyword 2",
        "SEO Keyword 3"
    ]
    
    # Create new DataFrame with selected columns
    new_format_df = all_results_df[selected_columns]
    
    # Save second Excel file with new format
    new_format_df.to_excel("All_Results_New_Format.xlsx", index=False, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}})
    print("Saved to All_Results_New_Format.xlsx")

if __name__ == "__main__":
    main()

main()



